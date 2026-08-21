from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(r"C:\Users\caaab\OneDrive\Documentos\Projects\abigobaldo\Abigobaldo\Documentos\Guia_de_Controles_Abigobaldo.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK = "1F4D78"
LIGHT = "E8EEF5"
INK = RGBColor(31, 77, 120)

def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)

def cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(3)
r = title.add_run("Abigobaldo: Guia Rápido de Controles")
set_font(r, 22, True, DARK)

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(14)
r = sub.add_run("Leia antes de começar. Olhe para os objetos e use os avisos que aparecem na HUD.")
set_font(r, 11, False, "555555")

intro = doc.add_paragraph()
intro.paragraph_format.space_after = Pt(10)
r = intro.add_run("Objetivo: ")
set_font(r, 11, True, DARK)
r = intro.add_run("atenda os pedidos dos clientes preparando e entregando a comida certa. Um pedido errado, ruim ou muito demorado reduz sua pontuação.")
set_font(r)

heading = doc.add_paragraph()
heading.paragraph_format.space_before = Pt(6)
heading.paragraph_format.space_after = Pt(7)
r = heading.add_run("CONTROLES PRINCIPAIS")
set_font(r, 16, True, BLUE)

table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
set_table_geometry(table, [2700, 6660])
headers = table.rows[0].cells
for cell, text in zip(headers, ("Controle", "O que faz")):
    shade(cell, LIGHT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, 11, True, DARK)

rows = [
    ("W, A, S, D", "Movem o Abigobaldo."),
    ("Mouse", "Move a câmera e aponta para objetos, clientes e estações."),
    ("E", "Interage com o que você está olhando: conversa, tira, deposita, usa ou ativa algo."),
    ("Clique esquerdo", "Pega objetos e usa certas estações. Segure para interações que exigem arrastar/manipular."),
    ("R + Mouse", "Gira o item que está na sua mão."),
    ("1, 2, 3 ou 4", "Escolhe uma resposta quando aparecerem opções de diálogo."),
    ("Enter", "Avança o diálogo quando não houver opções."),
    ("Esc", "Abre ou fecha o menu de pausa durante a partida."),
]
for key, action in rows:
    cells = table.add_row().cells
    cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)
    cells[1].paragraphs[0].paragraph_format.space_after = Pt(0)
    r = cells[0].paragraphs[0].add_run(key)
    set_font(r, 11, True, DARK)
    r = cells[1].paragraphs[0].add_run(action)
    set_font(r)

set_table_geometry(table, [2700, 6660])

heading = doc.add_paragraph()
heading.paragraph_format.space_before = Pt(16)
heading.paragraph_format.space_after = Pt(7)
r = heading.add_run("COMO ENTENDER A HUD")
set_font(r, 16, True, BLUE)

for label, description in [
    ("Avisos de interação", "só aparecem quando você está olhando para algo útil. Siga exatamente o botão indicado."),
    ("Pedido: (comida)", "aparece quando um cliente fez um pedido. Prepare essa comida e entregue ao cliente certo."),
    ("Exemplos", "olhando para um cliente: “Clique E para conversar”. Com comida na mão olhando para um recipiente: “Clique E para depositar”. Olhando para um spawner: “Clique Mouse Esquerdo para pegar”."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(label + ": ")
    set_font(r, 11, True, DARK)
    r = p.add_run(description)
    set_font(r)

note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(8)
note.paragraph_format.space_after = Pt(0)
r = note.add_run("Dica: ")
set_font(r, 11, True, "7A5A00")
r = note.add_run("os clientes só têm paciência por um tempo depois da conversa. Entregue rápido, mas entregue certo.")
set_font(r, 11, False, "7A5A00")

doc.save(OUT)
print(OUT)
